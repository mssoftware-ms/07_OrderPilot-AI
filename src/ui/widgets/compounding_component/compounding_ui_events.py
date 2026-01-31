"""compounding_ui_events.py
Event handlers, business logic, and data export for CompoundingPanel.

This mixin handles:
- User input events
- Calculation triggering (recompute logic)
- KPI and table updates
- Export to XLSX/DOCX/CSV
- Settings persistence
"""

from __future__ import annotations

from decimal import Decimal
from typing import Optional
from pathlib import Path

from PyQt6.QtCore import Qt, QSignalBlocker
from PyQt6.QtGui import QGuiApplication
from PyQt6.QtWidgets import QTableWidgetItem, QMessageBox, QFileDialog

from src.ui.widgets.compounding_component.calculator import (
    Params,
    simulate,
    solve_daily_profit_pct_for_target,
    to_csv_rows,
)

# Check optional export dependencies
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import WdTable_Alignment
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _to_dec(x: float) -> Decimal:
    """Convert float to Decimal."""
    return Decimal(str(x))


class CompoundingUIEventsMixin:
    """Mixin for event handling and user interactions."""

    def _on_user_edited(self, key: str) -> None:
        """Track which field the user last edited ('daily' or 'target')."""
        self._last_edited = key
        self.recompute(origin=f"user_{key}")

    def _get_params(self) -> Params:
        """Collect current parameters from UI widgets."""
        return Params(
            start_capital=_to_dec(self.sp_start_capital.value()),
            fee_open_pct=_to_dec(self.sp_fee_open_pct.value()),
            fee_close_pct=_to_dec(self.sp_fee_close_pct.value()),
            leverage=_to_dec(self.sp_leverage.value()),
            daily_profit_pct=_to_dec(self.sp_daily_profit_pct.value()),
            reinvest_pct=_to_dec(self.sp_reinvest_pct.value()),
            tax_pct=_to_dec(self.sp_tax_pct.value()),
            days=int(self.sp_days.value()),
            apply_losses_to_capital=bool(self.cb_apply_losses.isChecked()),
        )

    def recompute(self, origin: str = "") -> None:
        """
        Main calculation trigger.

        Handles:
        1. Parameter validation
        2. Solver logic (if target was edited)
        3. Simulation
        4. UI updates (KPIs, table, chart)

        Args:
            origin: Source of the recompute call (for debugging)
        """
        if self._suppress_updates:
            return

        # Validate parameters
        try:
            params = self._get_params()
            params.validate()
        except Exception as e:
            self.lbl_solver_status.setText(f"Eingaben ungültig: {e}")
            return

        # If user edited target, solve for daily profit percentage
        if self._last_edited == "target" and origin.startswith("user_"):
            target = _to_dec(self.sp_target_month_net.value())
            st = solve_daily_profit_pct_for_target(params, target_monthly_net=target)
            if st.ok and st.daily_profit_pct is not None:
                self._suppress_updates = True
                try:
                    with QSignalBlocker(self.sp_daily_profit_pct):
                        self.sp_daily_profit_pct.setValue(float(st.daily_profit_pct))
                finally:
                    self._suppress_updates = False
                self.lbl_solver_status.setText(
                    f"Solver OK (Iterationen: {st.iterations}). Monats-Netto: {st.achieved_monthly_net}€"
                )
                params = self._get_params()
            else:
                msg = st.message
                if st.achieved_monthly_net is not None:
                    msg += f" (Nächstes Ergebnis: {st.achieved_monthly_net}€)"
                self.lbl_solver_status.setText(msg)

        # Run simulation
        try:
            days, kpis = simulate(params)
        except Exception as e:
            self.lbl_solver_status.setText(f"Berechnung fehlgeschlagen: {e}")
            return

        # If user edited daily profit, update target
        if self._last_edited == "daily" and origin.startswith("user_"):
            self._suppress_updates = True
            try:
                with QSignalBlocker(self.sp_target_month_net):
                    self.sp_target_month_net.setValue(float(kpis.sum_net))
            finally:
                self._suppress_updates = False
            self.lbl_solver_status.setText("Ziel netto im Monat aus Tagesgewinn% berechnet.")

        # Update UI
        self._update_kpis(kpis)
        self._update_table(days)
        self._update_plot(days)

    def _update_kpis(self, k) -> None:
        """Update KPI labels with calculation results."""
        self.kpi_end_capital.setText(f"{k.end_capital:.2f} €")
        self.kpi_sum_gross.setText(f"{k.sum_gross:.2f} €")
        self.kpi_sum_fees.setText(f"{k.sum_fees:.2f} €")
        self.kpi_sum_taxes.setText(f"{k.sum_taxes:.2f} €")
        self.kpi_sum_net.setText(f"{k.sum_net:.2f} €")
        self.kpi_sum_reinvest.setText(f"{k.sum_reinvest:.2f} €")
        self.kpi_sum_withdrawal.setText(f"{k.sum_withdrawal:.2f} €")
        self.kpi_roi.setText(f"{k.roi_pct} %")

    def _update_table(self, days) -> None:
        """Update data table (hidden but used for export)."""
        self.table.setRowCount(len(days))
        for r_i, r in enumerate(days):
            values = [
                r.day, r.start_capital, r.gross_profit, r.notional, r.fee_amount,
                r.profit_before_tax, r.tax_amount, r.net_profit, r.reinvest,
                r.withdrawal, r.end_capital
            ]
            for c_i, v in enumerate(values):
                item = QTableWidgetItem(str(v) if c_i == 0 else f"{v:.2f}")
                item.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                self.table.setItem(r_i, c_i, item)

    def _export_xlsx(self) -> None:
        """Export data to styled Excel file."""
        if not OPENPYXL_AVAILABLE:
            QMessageBox.warning(
                self, "Export nicht verfügbar",
                "openpyxl ist nicht installiert. Bitte 'pip install openpyxl' ausführen."
            )
            return

        try:
            params = self._get_params()
            days, kpis = simulate(params)
        except Exception as e:
            QMessageBox.critical(self, "Export fehlgeschlagen", str(e))
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "XLSX speichern", "compounding_report.xlsx", "Excel (*.xlsx)"
        )
        if not path:
            return

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Compounding Daten"

            # Header styling
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="2C5F2C", end_color="2C5F2C", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Headers
            headers = [
                "Tag", "Startkapital", "Brutto", "Notional", "Gebühren",
                "Vor Steuern", "Steuern", "Netto", "Reinvest", "Entnahme", "Endkapital"
            ]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = thin_border

            # Data rows
            money_format = '#,##0.00 €'
            for row_idx, d in enumerate(days, 2):
                values = [
                    d.day, float(d.start_capital), float(d.gross_profit), float(d.notional),
                    float(d.fee_amount), float(d.profit_before_tax), float(d.tax_amount),
                    float(d.net_profit), float(d.reinvest), float(d.withdrawal), float(d.end_capital)
                ]
                for col_idx, value in enumerate(values, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = thin_border
                    if col_idx > 1:  # Money columns
                        cell.number_format = money_format
                        cell.alignment = Alignment(horizontal="right")

            # KPIs summary sheet
            ws_kpi = wb.create_sheet("KPIs")
            kpi_data = [
                ("Endkapital", float(kpis.end_capital)),
                ("Summe Brutto", float(kpis.sum_gross)),
                ("Summe Gebühren", float(kpis.sum_fees)),
                ("Summe Steuern", float(kpis.sum_taxes)),
                ("Summe Netto", float(kpis.sum_net)),
                ("Summe reinvestiert", float(kpis.sum_reinvest)),
                ("Summe entnommen", float(kpis.sum_withdrawal)),
                ("ROI", f"{kpis.roi_pct}%"),
            ]
            for row_idx, (label, value) in enumerate(kpi_data, 1):
                ws_kpi.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
                cell = ws_kpi.cell(row=row_idx, column=2, value=value)
                if isinstance(value, float):
                    cell.number_format = money_format

            # Auto-adjust column widths
            for ws_sheet in [ws, ws_kpi]:
                for column in ws_sheet.columns:
                    max_length = max(len(str(cell.value or "")) for cell in column)
                    ws_sheet.column_dimensions[column[0].column_letter].width = min(max_length + 2, 20)

            wb.save(path)
            QMessageBox.information(self, "Export erfolgreich", f"Daten wurden nach {path} exportiert.")

        except Exception as e:
            QMessageBox.critical(self, "XLSX Export Fehler", str(e))

    def _export_docx(self) -> None:
        """Export data to Word document with formatted table."""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(
                self, "Export nicht verfügbar",
                "python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen."
            )
            return

        try:
            params = self._get_params()
            days, kpis = simulate(params)
        except Exception as e:
            QMessageBox.critical(self, "Export fehlgeschlagen", str(e))
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "DOCX speichern", "compounding_report.docx", "Word (*.docx)"
        )
        if not path:
            return

        try:
            doc = Document()

            # Title
            doc.add_heading('Compounding & P&L Report', 0)

            # KPIs section
            doc.add_heading('Monatliche Kennzahlen (KPIs)', level=1)
            kpi_table = doc.add_table(rows=8, cols=2)
            kpi_table.style = 'Table Grid'
            kpi_data = [
                ("Endkapital", f"{kpis.end_capital:.2f} €"),
                ("Summe Brutto", f"{kpis.sum_gross:.2f} €"),
                ("Summe Gebühren", f"{kpis.sum_fees:.2f} €"),
                ("Summe Steuern", f"{kpis.sum_taxes:.2f} €"),
                ("Summe Netto", f"{kpis.sum_net:.2f} €"),
                ("Summe reinvestiert", f"{kpis.sum_reinvest:.2f} €"),
                ("Summe entnommen", f"{kpis.sum_withdrawal:.2f} €"),
                ("ROI", f"{kpis.roi_pct}%"),
            ]
            for i, (label, value) in enumerate(kpi_data):
                row = kpi_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Bold the labels
                row.cells[0].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()  # Spacing

            # Daily data section
            doc.add_heading('Tägliche Übersicht', level=1)

            headers = ["Tag", "Start", "Brutto", "Gebühren", "Steuern", "Netto", "Endkapital"]
            data_table = doc.add_table(rows=len(days) + 1, cols=len(headers))
            data_table.style = 'Table Grid'

            # Headers
            header_row = data_table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                header_row.cells[i].paragraphs[0].runs[0].bold = True

            # Data
            for row_idx, d in enumerate(days, 1):
                row = data_table.rows[row_idx]
                row.cells[0].text = str(d.day)
                row.cells[1].text = f"{d.start_capital:.2f}"
                row.cells[2].text = f"{d.gross_profit:.2f}"
                row.cells[3].text = f"{d.fee_amount:.2f}"
                row.cells[4].text = f"{d.tax_amount:.2f}"
                row.cells[5].text = f"{d.net_profit:.2f}"
                row.cells[6].text = f"{d.end_capital:.2f}"

            doc.save(path)
            QMessageBox.information(self, "Export erfolgreich", f"Report wurde nach {path} exportiert.")

        except Exception as e:
            QMessageBox.critical(self, "DOCX Export Fehler", str(e))

    def _copy_table_to_clipboard(self) -> None:
        """Copy table data to clipboard as CSV."""
        try:
            params = self._get_params()
            days, _k = simulate(params)
            rows = to_csv_rows(days)
            text = "\n".join(["\t".join(r) for r in rows])
        except Exception as e:
            QMessageBox.critical(self, "Kopieren fehlgeschlagen", str(e))
            return
        QGuiApplication.clipboard().setText(text)
        QMessageBox.information(self, "Kopiert", "Tages-Tabelle wurde in die Zwischenablage kopiert.")

    # ==================== Settings Persistence ====================

    def _on_tab_changed(self, index: int) -> None:
        """Save settings when leaving Compounding tab."""
        # Save settings whenever tab changes (covers leaving the inputs tab)
        self._save_settings()

    def _save_settings(self) -> None:
        """Save current input values to QSettings."""
        try:
            self._settings.setValue("start_capital", self.sp_start_capital.value())
            self._settings.setValue("vip_level", self.cb_vip_level.currentText())
            self._settings.setValue("order_type", self.cb_order_type.currentText())
            self._settings.setValue("fee_open_pct", self.sp_fee_open_pct.value())
            self._settings.setValue("fee_close_pct", self.sp_fee_close_pct.value())
            self._settings.setValue("leverage", self.sp_leverage.value())
            self._settings.setValue("daily_profit_pct", self.sp_daily_profit_pct.value())
            self._settings.setValue("reinvest_pct", self.sp_reinvest_pct.value())
            self._settings.setValue("tax_pct", self.sp_tax_pct.value())
            self._settings.setValue("days", self.sp_days.value())
            self._settings.setValue("trading_days_preset", self.cb_trading_days_preset.currentIndex())
            self._settings.setValue("apply_losses", self.cb_apply_losses.isChecked())
            self._settings.setValue("target_month_net", self.sp_target_month_net.value())
            self._settings.setValue("plot_mode", self.cb_plot_mode.currentIndex())
            self._settings.sync()
        except Exception:
            pass  # Silently ignore save errors

    def _load_settings(self) -> None:
        """Load saved input values from QSettings."""
        self._suppress_updates = True
        try:
            if self._settings.contains("start_capital"):
                self.sp_start_capital.setValue(float(self._settings.value("start_capital", 100.0)))
            if self._settings.contains("vip_level"):
                self.cb_vip_level.setCurrentText(str(self._settings.value("vip_level", "VIP0")))
            if self._settings.contains("order_type"):
                self.cb_order_type.setCurrentText(str(self._settings.value("order_type", "Taker/Taker")))
            if self._settings.contains("fee_open_pct"):
                self.sp_fee_open_pct.setValue(float(self._settings.value("fee_open_pct", 0.06)))
            if self._settings.contains("fee_close_pct"):
                self.sp_fee_close_pct.setValue(float(self._settings.value("fee_close_pct", 0.06)))
            if self._settings.contains("leverage"):
                self.sp_leverage.setValue(float(self._settings.value("leverage", 20.0)))
            if self._settings.contains("daily_profit_pct"):
                self.sp_daily_profit_pct.setValue(float(self._settings.value("daily_profit_pct", 30.0)))
            if self._settings.contains("reinvest_pct"):
                self.sp_reinvest_pct.setValue(float(self._settings.value("reinvest_pct", 30.0)))
            if self._settings.contains("tax_pct"):
                self.sp_tax_pct.setValue(float(self._settings.value("tax_pct", 25.0)))
            if self._settings.contains("days"):
                self.sp_days.setValue(int(self._settings.value("days", 30)))
            if self._settings.contains("trading_days_preset"):
                self.cb_trading_days_preset.setCurrentIndex(int(self._settings.value("trading_days_preset", 0)))
            if self._settings.contains("apply_losses"):
                self.cb_apply_losses.setChecked(self._settings.value("apply_losses", False) in [True, "true", "True"])
            if self._settings.contains("target_month_net"):
                self.sp_target_month_net.setValue(float(self._settings.value("target_month_net", 0.0)))
            if self._settings.contains("plot_mode"):
                self.cb_plot_mode.setCurrentIndex(int(self._settings.value("plot_mode", 0)))
        except Exception:
            pass  # Use defaults on load errors
        finally:
            self._suppress_updates = False

    def closeEvent(self, event) -> None:
        """Save settings when widget is closed."""
        self._save_settings()
        super().closeEvent(event)
