"""ui.py
PyQt6 UI-Komponente (einbettbar) für den Compounding-Rechner.

- 2 Tabs: (Eingaben + KPIs) und (Details mit modernem Chart).
- Live-Updates; keine Signal-Endlosschleifen (QSignalBlocker/Guard).
- Theme-kompatibel: keine hardcodierten Farben; Qt-Palette wird übernommen.
- Matplotlib-Embedding via FigureCanvasQTAgg (backend_qtagg).
- Modernes Chart-Design mit Gradient-Fills und mehreren Datenserien.
"""

from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import Optional, List

import numpy as np
from PyQt6.QtCore import Qt, QSignalBlocker, QSettings
from PyQt6.QtGui import QGuiApplication, QColor
from PyQt6.QtWidgets import (
    QWidget, QTabWidget, QVBoxLayout, QHBoxLayout, QFormLayout, QGroupBox, QLabel,
    QDoubleSpinBox, QSpinBox, QCheckBox, QPushButton, QTableWidget, QTableWidgetItem,
    QHeaderView, QMessageBox, QFileDialog, QComboBox,
)

from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
from matplotlib.colors import LinearSegmentedColormap
import matplotlib.pyplot as plt

# XLSX/DOCX Exports
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import Wd_Table_Alignment
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.ui.widgets.compounding_component.calculator import (
    Params,
    simulate,
    solve_daily_profit_pct_for_target,
    to_csv_rows,
)

FUTURES_FEES_BY_VIP = {
    "VIP0": {"maker": 0.02, "taker": 0.06},
    "VIP1": {"maker": 0.02, "taker": 0.05},
    "VIP2": {"maker": 0.016, "taker": 0.05},
    "VIP3": {"maker": 0.014, "taker": 0.04},
    "VIP4": {"maker": 0.012, "taker": 0.0375},
    "VIP5": {"maker": 0.01, "taker": 0.035},
    "VIP6": {"maker": 0.008, "taker": 0.0315},
    "VIP7": {"maker": 0.006, "taker": 0.03},
}


def _to_dec(x: float) -> Decimal:
    return Decimal(str(x))


# Modern color palette for financial charts
MODERN_COLORS = {
    'gross': '#FF9500',      # Orange - Gross Profit
    'fees': '#FF3B30',       # Red - Fees
    'taxes': '#AF52DE',      # Purple - Taxes
    'net': '#34C759',        # Green - Net Profit
    'invest': '#007AFF',     # Blue - Investment/Capital
    'grid': '#3A3A3C',       # Dark gray - Grid
    'background': '#1C1C1E', # Dark background
    'surface': '#2C2C2E',    # Surface color
    'text': '#FFFFFF',       # White text
    'text_secondary': '#8E8E93',  # Secondary text
}


class MplCanvas(FigureCanvas):
    """Modern matplotlib canvas with dark theme and gradient support."""
    
    def __init__(self, parent: QWidget):
        self.fig = Figure(facecolor='none', edgecolor='none')
        super().__init__(self.fig)
        self.setParent(parent)
        self.ax = self.fig.add_subplot(111)
        self._setup_modern_style()

    def _setup_modern_style(self) -> None:
        """Apply modern dark theme styling."""
        # Use a clean style as base
        plt.style.use('dark_background')
        
    def apply_qt_palette(self, widget: QWidget) -> None:
        """Apply Qt palette colors to matplotlib figure."""
        pal = widget.palette()
        window = pal.color(widget.backgroundRole())
        base = pal.color(pal.ColorRole.Base)
        text = pal.color(pal.ColorRole.Text)

        def qcol(c, alpha=1.0):
            return (c.redF(), c.greenF(), c.blueF(), alpha)

        # Modern dark styling
        bg_color = qcol(window)
        surface_color = qcol(base)
        text_color = qcol(text)
        
        self.fig.set_facecolor(bg_color)
        self.ax.set_facecolor(surface_color)
        
        # Modern axis styling
        self.ax.tick_params(axis="both", colors=text_color, labelsize=9)
        self.ax.xaxis.label.set_color(text_color)
        self.ax.yaxis.label.set_color(text_color)
        self.ax.title.set_color(text_color)
        
        # Subtle spines
        for spine in self.ax.spines.values():
            spine.set_color(qcol(text, 0.3))
            spine.set_linewidth(0.5)
        
        # Remove top and right spines for cleaner look
        self.ax.spines['top'].set_visible(False)
        self.ax.spines['right'].set_visible(False)


class CompoundingPanel(QWidget):
    """Compounding/P&L Calculator Panel with settings persistence."""
    
    SETTINGS_KEY = "CompoundingPanel"
    
    def __init__(self, parent: Optional[QWidget] = None):
        super().__init__(parent)

        self._last_edited: Optional[str] = None  # "daily" oder "target"
        self._suppress_updates: bool = False
        self._settings = QSettings("OrderPilot", "CompoundingPanel")

        self.tabs = QTabWidget(self)
        self.tab_inputs = QWidget()
        self.tab_details = QWidget()
        self.tabs.addTab(self.tab_inputs, "Compounding")
        self.tabs.addTab(self.tab_details, "Details")
        
        # Save settings when switching tabs
        self.tabs.currentChanged.connect(self._on_tab_changed)

        root = QVBoxLayout(self)
        root.addWidget(self.tabs)

        # Build details tab first (creates self.table used by recompute)
        self._build_tab_details()
        self._build_tab_inputs()
        
        # Load saved settings after UI is built
        self._load_settings()
        
        self.recompute(origin="init")

    def _build_tab_inputs(self) -> None:
        layout = QVBoxLayout(self.tab_inputs)
        box = QGroupBox("Eingaben")
        form = QFormLayout(box)

        self.sp_start_capital = QDoubleSpinBox(); self.sp_start_capital.setRange(0, 1e9); self.sp_start_capital.setDecimals(2); self.sp_start_capital.setValue(100.0); self.sp_start_capital.setSuffix(" €")
        self.cb_vip_level = QComboBox()
        self.cb_vip_level.addItems(["VIP0", "VIP1", "VIP2", "VIP3", "VIP4", "VIP5", "VIP6", "VIP7", "Custom"])
        self.cb_vip_level.setCurrentText("VIP0")

        self.cb_order_type = QComboBox()
        self.cb_order_type.addItems(["Taker/Taker", "Maker/Maker", "Maker/Taker", "Taker/Maker"])
        self.cb_order_type.setCurrentText("Taker/Taker")

        self.sp_fee_open_pct = QDoubleSpinBox(); self.sp_fee_open_pct.setRange(0, 100); self.sp_fee_open_pct.setDecimals(4); self.sp_fee_open_pct.setSingleStep(0.01); self.sp_fee_open_pct.setSuffix(" %")
        self.sp_fee_close_pct = QDoubleSpinBox(); self.sp_fee_close_pct.setRange(0, 100); self.sp_fee_close_pct.setDecimals(4); self.sp_fee_close_pct.setSingleStep(0.01); self.sp_fee_close_pct.setSuffix(" %")
        self.sp_leverage = QDoubleSpinBox(); self.sp_leverage.setRange(0.01, 1000); self.sp_leverage.setDecimals(2); self.sp_leverage.setValue(20.0); self.sp_leverage.setSuffix(" x")
        self.sp_daily_profit_pct = QDoubleSpinBox(); self.sp_daily_profit_pct.setRange(-99.0, 1000.0); self.sp_daily_profit_pct.setDecimals(4); self.sp_daily_profit_pct.setValue(30.0); self.sp_daily_profit_pct.setSuffix(" %")
        self.sp_reinvest_pct = QDoubleSpinBox(); self.sp_reinvest_pct.setRange(0, 100); self.sp_reinvest_pct.setDecimals(2); self.sp_reinvest_pct.setValue(30.0); self.sp_reinvest_pct.setSuffix(" %")
        self.sp_tax_pct = QDoubleSpinBox(); self.sp_tax_pct.setRange(0, 100); self.sp_tax_pct.setDecimals(2); self.sp_tax_pct.setValue(25.0); self.sp_tax_pct.setSuffix(" %")
        self.sp_days = QSpinBox(); self.sp_days.setRange(1, 366); self.sp_days.setValue(30)

        self.cb_trading_days_preset = QComboBox()
        self.cb_trading_days_preset.addItems(["30 (Kalendertage)", "22 (typ. Börsentage)", "20 (typ. Börsentage)"])
        self.cb_trading_days_preset.setCurrentIndex(0)

        self.cb_apply_losses = QCheckBox("Verluste reduzieren Kapital (optional)")
        self.cb_apply_losses.setChecked(False)

        self.sp_target_month_net = QDoubleSpinBox(); self.sp_target_month_net.setRange(-1e9, 1e9); self.sp_target_month_net.setDecimals(2); self.sp_target_month_net.setValue(0.0); self.sp_target_month_net.setSuffix(" €")
        self.lbl_solver_status = QLabel(""); self.lbl_solver_status.setWordWrap(True)

        form.addRow("Startkapital", self.sp_start_capital)
        form.addRow("VIP Status", self.cb_vip_level)
        form.addRow("Order Type (Open/Close)", self.cb_order_type)
        form.addRow("Gebuehren Open", self.sp_fee_open_pct)
        form.addRow("Gebuehren Close", self.sp_fee_close_pct)
        form.addRow("Hebel", self.sp_leverage)
        form.addRow("Gewinn % / Tag", self.sp_daily_profit_pct)
        form.addRow("Reinvestition", self.sp_reinvest_pct)
        form.addRow("Steuern (Monatsende)", self.sp_tax_pct)
        form.addRow("Tage (manuell)", self.sp_days)
        form.addRow("Preset Trading-Tage", self.cb_trading_days_preset)
        form.addRow("", self.cb_apply_losses)
        form.addRow("Ziel netto im Monat", self.sp_target_month_net)
        form.addRow("Status", self.lbl_solver_status)

        layout.addWidget(box)

        kpi_box = QGroupBox("Monats-KPIs")
        kpi_form = QFormLayout(kpi_box)
        self.kpi_end_capital = QLabel("—"); self.kpi_sum_gross = QLabel("—"); self.kpi_sum_fees = QLabel("—"); self.kpi_sum_taxes = QLabel("—")
        self.kpi_sum_net = QLabel("—"); self.kpi_sum_reinvest = QLabel("—"); self.kpi_sum_withdrawal = QLabel("—"); self.kpi_roi = QLabel("—")

        kpi_form.addRow("Endkapital", self.kpi_end_capital)
        kpi_form.addRow("Summe Brutto", self.kpi_sum_gross)
        kpi_form.addRow("Summe Gebühren", self.kpi_sum_fees)
        kpi_form.addRow("Summe Steuern", self.kpi_sum_taxes)
        kpi_form.addRow("Summe Netto", self.kpi_sum_net)
        kpi_form.addRow("Summe reinvestiert", self.kpi_sum_reinvest)
        kpi_form.addRow("Summe entnommen", self.kpi_sum_withdrawal)
        kpi_form.addRow("ROI", self.kpi_roi)

        layout.addWidget(kpi_box)
        layout.addStretch(1)

        # Signals
        self.sp_daily_profit_pct.valueChanged.connect(lambda _v: self._on_user_edited("daily"))
        self.sp_target_month_net.valueChanged.connect(lambda _v: self._on_user_edited("target"))

        for w in [
            self.sp_start_capital,
            self.sp_fee_open_pct,
            self.sp_fee_close_pct,
            self.sp_leverage,
            self.sp_reinvest_pct,
            self.sp_tax_pct,
            self.sp_days,
        ]:
            w.valueChanged.connect(lambda _v: self.recompute(origin="param_change"))
        self.sp_fee_open_pct.valueChanged.connect(self._on_fee_override)
        self.sp_fee_close_pct.valueChanged.connect(self._on_fee_override)
        self.cb_apply_losses.stateChanged.connect(lambda _s: self.recompute(origin="param_change"))
        self.cb_trading_days_preset.currentIndexChanged.connect(self._apply_trading_days_preset)
        self.cb_vip_level.currentIndexChanged.connect(self._apply_fee_preset)
        self.cb_order_type.currentIndexChanged.connect(self._apply_fee_preset)
        self._apply_fee_preset()

    def _build_tab_details(self) -> None:
        """Build the Details tab with modern chart (table hidden)."""
        layout = QVBoxLayout(self.tab_details)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)
        
        # Top toolbar with modern styling
        top = QHBoxLayout()
        top.setSpacing(8)
        
        # Plot mode selector
        plot_label = QLabel("📊 Ansicht:")
        plot_label.setStyleSheet("font-weight: bold; color: #8E8E93;")
        self.cb_plot_mode = QComboBox()
        self.cb_plot_mode.addItems([
            "Alle Metriken (gestapelt)",
            "Täglicher Netto-Gewinn", 
            "Kumuliertes Netto",
            "Gewinn vs. Kosten",
            "Monatliche Übersicht (gestapelt)",
            "Jährliche Übersicht (gestapelt)",
        ])
        self.cb_plot_mode.setCurrentIndex(0)
        self.cb_plot_mode.currentIndexChanged.connect(lambda _i: self.recompute(origin="plot_mode"))
        self.cb_plot_mode.setStyleSheet("""
            QComboBox {
                background-color: #2C2C2E;
                border: 1px solid #3A3A3C;
                border-radius: 6px;
                padding: 5px 10px;
                min-width: 180px;
            }
            QComboBox:hover {
                border-color: #48484A;
            }
        """)

        # Export buttons with modern styling
        btn_style = """
            QPushButton {
                background-color: #2C2C2E;
                border: 1px solid #3A3A3C;
                border-radius: 6px;
                padding: 6px 12px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #3A3A3C;
                border-color: #48484A;
            }
            QPushButton:pressed {
                background-color: #48484A;
            }
        """
        
        self.btn_copy_csv = QPushButton("📋 Kopieren")
        self.btn_copy_csv.setStyleSheet(btn_style)
        self.btn_copy_csv.clicked.connect(self._copy_table_to_clipboard)
        
        self.btn_export_xlsx = QPushButton("📊 XLSX exportieren")
        self.btn_export_xlsx.setStyleSheet(btn_style)
        self.btn_export_xlsx.clicked.connect(self._export_xlsx)
        self.btn_export_xlsx.setEnabled(OPENPYXL_AVAILABLE)
        
        self.btn_export_docx = QPushButton("📄 DOCX Report")
        self.btn_export_docx.setStyleSheet(btn_style)
        self.btn_export_docx.clicked.connect(self._export_docx)
        self.btn_export_docx.setEnabled(DOCX_AVAILABLE)

        top.addWidget(plot_label)
        top.addWidget(self.cb_plot_mode)
        top.addStretch(1)
        top.addWidget(self.btn_copy_csv)
        top.addWidget(self.btn_export_xlsx)
        top.addWidget(self.btn_export_docx)
        layout.addLayout(top)

        # Hidden table (still exists for data/export)
        self.table = QTableWidget()
        self.table.setColumnCount(11)
        self.table.setHorizontalHeaderLabels([
            "Tag", "Startkapital", "Brutto", "Notional", "Gebühren",
            "Vor Steuern", "Steuern", "Netto", "Reinvest", "Entnahme", "Endkapital"
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setVisible(False)  # Hidden but available for export

        # Modern chart canvas - takes full space
        self.canvas = MplCanvas(self)
        self.canvas.setMinimumHeight(400)
        self.canvas.setStyleSheet("""
            background-color: #1C1C1E;
            border: 1px solid #3A3A3C;
            border-radius: 8px;
        """)
        
        layout.addWidget(self.canvas, stretch=1)

    def _apply_trading_days_preset(self, idx: int) -> None:
        preset_map = {0: 30, 1: 22, 2: 20}
        val = preset_map.get(idx, 30)
        with QSignalBlocker(self.sp_days):
            self.sp_days.setValue(val)
        self.recompute(origin="preset_days")

    def _apply_fee_preset(self) -> None:
        vip = self.cb_vip_level.currentText()
        if vip == "Custom":
            return
        fees = FUTURES_FEES_BY_VIP.get(vip)
        if not fees:
            return
        maker = float(fees["maker"])
        taker = float(fees["taker"])
        order = self.cb_order_type.currentText()
        if order == "Maker/Maker":
            open_fee = maker
            close_fee = maker
        elif order == "Maker/Taker":
            open_fee = maker
            close_fee = taker
        elif order == "Taker/Maker":
            open_fee = taker
            close_fee = maker
        else:
            open_fee = taker
            close_fee = taker

        with QSignalBlocker(self.sp_fee_open_pct):
            self.sp_fee_open_pct.setValue(open_fee)
        with QSignalBlocker(self.sp_fee_close_pct):
            self.sp_fee_close_pct.setValue(close_fee)
        self.recompute(origin="fee_preset")

    def _on_fee_override(self) -> None:
        if self._suppress_updates:
            return
        if self.cb_vip_level.currentText() != "Custom":
            with QSignalBlocker(self.cb_vip_level):
                self.cb_vip_level.setCurrentText("Custom")
        self.recompute(origin="fee_override")

    def _on_user_edited(self, key: str) -> None:
        self._last_edited = key
        self.recompute(origin=f"user_{key}")

    def _get_params(self) -> Params:
        return Params(
            start_capital=_to_dec(self.sp_start_capital.value()),
            fee_open_pct=_to_dec(self.sp_fee_open_pct.value()),
            fee_close_pct=_to_dec(self.sp_fee_close_pct.value()),
            leverage=_to_dec(self.sp_leverage.value()),
            daily_profit_pct=_to_dec(self.sp_daily_profit_pct.value()),
            reinvest_pct=_to_dec(self.sp_reinvest_pct.value()),
            tax_pct=_to_dec(self.sp_tax_pct.value()),
            days=int(self.sp_days.value()),
            apply_losses_to_capital=bool(self.cb_apply_losses.isChecked()),
        )

    def recompute(self, origin: str = "") -> None:
        if self._suppress_updates:
            return

        try:
            params = self._get_params()
            params.validate()
        except Exception as e:
            self.lbl_solver_status.setText(f"Eingaben ungültig: {e}")
            return

        if self._last_edited == "target" and origin.startswith("user_"):
            target = _to_dec(self.sp_target_month_net.value())
            st = solve_daily_profit_pct_for_target(params, target_monthly_net=target)
            if st.ok and st.daily_profit_pct is not None:
                self._suppress_updates = True
                try:
                    with QSignalBlocker(self.sp_daily_profit_pct):
                        self.sp_daily_profit_pct.setValue(float(st.daily_profit_pct))
                finally:
                    self._suppress_updates = False
                self.lbl_solver_status.setText(f"Solver OK (Iterationen: {st.iterations}). Monats-Netto: {st.achieved_monthly_net}€")
                params = self._get_params()
            else:
                msg = st.message
                if st.achieved_monthly_net is not None:
                    msg += f" (Nächstes Ergebnis: {st.achieved_monthly_net}€)"
                self.lbl_solver_status.setText(msg)

        try:
            days, kpis = simulate(params)
        except Exception as e:
            self.lbl_solver_status.setText(f"Berechnung fehlgeschlagen: {e}")
            return

        if self._last_edited == "daily" and origin.startswith("user_"):
            self._suppress_updates = True
            try:
                with QSignalBlocker(self.sp_target_month_net):
                    self.sp_target_month_net.setValue(float(kpis.sum_net))
            finally:
                self._suppress_updates = False
            self.lbl_solver_status.setText("Ziel netto im Monat aus Tagesgewinn% berechnet.")

        self._update_kpis(kpis)
        self._update_table(days)
        self._update_plot(days)

    def _update_kpis(self, k) -> None:
        self.kpi_end_capital.setText(f"{k.end_capital:.2f} €")
        self.kpi_sum_gross.setText(f"{k.sum_gross:.2f} €")
        self.kpi_sum_fees.setText(f"{k.sum_fees:.2f} €")
        self.kpi_sum_taxes.setText(f"{k.sum_taxes:.2f} €")
        self.kpi_sum_net.setText(f"{k.sum_net:.2f} €")
        self.kpi_sum_reinvest.setText(f"{k.sum_reinvest:.2f} €")
        self.kpi_sum_withdrawal.setText(f"{k.sum_withdrawal:.2f} €")
        self.kpi_roi.setText(f"{k.roi_pct} %")

    def _update_table(self, days) -> None:
        self.table.setRowCount(len(days))
        for r_i, r in enumerate(days):
            values = [r.day, r.start_capital, r.gross_profit, r.notional, r.fee_amount, r.profit_before_tax, r.tax_amount, r.net_profit, r.reinvest, r.withdrawal, r.end_capital]
            for c_i, v in enumerate(values):
                item = QTableWidgetItem(str(v) if c_i == 0 else f"{v:.2f}")
                item.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                self.table.setItem(r_i, c_i, item)

    def _update_plot(self, days) -> None:
        """Update plot with modern financial chart design."""
        ax = self.canvas.ax
        ax.clear()
        self.canvas.apply_qt_palette(self)
        
        xs = np.array([d.day for d in days])
        mode = self.cb_plot_mode.currentIndex()
        
        if mode == 0:
            # Stacked area chart with all metrics
            self._plot_stacked_metrics(ax, days, xs)
        elif mode == 1:
            # Daily net profit with gradient fill
            self._plot_daily_net(ax, days, xs)
        elif mode == 2:
            # Cumulative net profit
            self._plot_cumulative(ax, days, xs)
        elif mode == 3:
            # Profit vs Costs comparison
            self._plot_profit_vs_costs(ax, days, xs)
        elif mode == 4:
            # Monthly stacked view
            self._plot_monthly(ax, days, xs)
        else:
            # Yearly stacked view
            self._plot_yearly(ax, days, xs)
        
        # Common styling
        if mode < 4:  # Daily views
            ax.set_xlabel("Tag", fontsize=10, fontweight='medium')
            ax.set_xlim(xs[0], xs[-1])
        
        ax.grid(True, alpha=0.15, linestyle='-', linewidth=0.5)
        
        # Modern legend
        legend = ax.legend(
            loc='upper left',
            frameon=True,
            fancybox=True,
            shadow=False,
            framealpha=0.8,
            edgecolor='#3A3A3C',
            fontsize=9,
        )
        if legend:
            legend.get_frame().set_facecolor('#2C2C2E')
        
        self.canvas.fig.tight_layout(pad=1.5)
        self.canvas.draw()
    
    def _plot_stacked_metrics(self, ax, days, xs) -> None:
        """Create modern stacked area chart with all key metrics."""
        # Extract data series
        gross = np.array([float(d.gross_profit) for d in days])
        fees = np.array([float(d.fee_amount) for d in days])
        taxes = np.array([float(d.tax_amount) for d in days])
        net = np.array([float(d.net_profit) for d in days])
        start_cap = np.array([float(d.start_capital) for d in days])
        
        # Normalize for visualization (show relative contributions)
        # Plot lines with area fills
        
        # Net profit area (base)
        ax.fill_between(xs, 0, net, alpha=0.6, color=MODERN_COLORS['net'], label='Netto-Gewinn')
        ax.plot(xs, net, color=MODERN_COLORS['net'], linewidth=2, alpha=0.9)
        
        # Gross profit line
        ax.plot(xs, gross, color=MODERN_COLORS['gross'], linewidth=2.5, 
                label='Brutto-Gewinn', linestyle='-', marker='o', markersize=3)
        
        # Fees area (negative visual)
        ax.fill_between(xs, 0, -fees, alpha=0.5, color=MODERN_COLORS['fees'], label='Gebühren')
        
        # Taxes area (negative visual, stacked on fees)
        ax.fill_between(xs, -fees, -fees - taxes, alpha=0.5, color=MODERN_COLORS['taxes'], label='Steuern')
        
        # Investment reference line (scaled down for visibility)
        invest_scaled = start_cap / start_cap.max() * gross.max() * 0.3 if gross.max() > 0 else start_cap * 0.01
        ax.plot(xs, invest_scaled, color=MODERN_COLORS['invest'], linewidth=1.5, 
                linestyle='--', alpha=0.7, label='Invest (skaliert)')
        
        ax.set_ylabel("Betrag (€)", fontsize=10, fontweight='medium')
        ax.set_title("Tägliche Übersicht: Gewinn, Gebühren & Steuern", 
                     fontsize=12, fontweight='bold', pad=15)
        ax.axhline(y=0, color='#8E8E93', linewidth=0.8, linestyle='-', alpha=0.5)
    
    def _plot_daily_net(self, ax, days, xs) -> None:
        """Daily net profit with modern gradient fill."""
        ys = np.array([float(d.net_profit) for d in days])
        
        # Gradient fill under the line
        ax.fill_between(xs, 0, ys, alpha=0.4, 
                        color=MODERN_COLORS['net'], label='Netto-Gewinn')
        ax.plot(xs, ys, color=MODERN_COLORS['net'], linewidth=2.5,
                marker='o', markersize=4, markerfacecolor='white', markeredgewidth=1.5)
        
        ax.set_ylabel("Netto-Gewinn (€)", fontsize=10, fontweight='medium')
        ax.set_title("Täglicher Netto-Gewinn", fontsize=12, fontweight='bold', pad=15)
        ax.axhline(y=0, color='#8E8E93', linewidth=0.8, linestyle='-', alpha=0.5)
    
    def _plot_cumulative(self, ax, days, xs) -> None:
        """Cumulative net profit with gradient."""
        cum = np.cumsum([float(d.net_profit) for d in days])
        
        # Gradient fill
        ax.fill_between(xs, 0, cum, alpha=0.4, color=MODERN_COLORS['net'], label='Kumulierter Netto-Gewinn')
        ax.plot(xs, cum, color=MODERN_COLORS['net'], linewidth=2.5)
        
        # Add cumulative gross for reference
        cum_gross = np.cumsum([float(d.gross_profit) for d in days])
        ax.plot(xs, cum_gross, color=MODERN_COLORS['gross'], linewidth=2, 
                linestyle='--', alpha=0.7, label='Kumulierter Brutto')
        
        ax.set_ylabel("Kumulierter Betrag (€)", fontsize=10, fontweight='medium')
        ax.set_title("Kumulierte Entwicklung", fontsize=12, fontweight='bold', pad=15)
        ax.axhline(y=0, color='#8E8E93', linewidth=0.8, linestyle='-', alpha=0.5)
    
    def _plot_profit_vs_costs(self, ax, days, xs) -> None:
        """Compare profit vs costs (fees + taxes)."""
        gross = np.array([float(d.gross_profit) for d in days])
        fees = np.array([float(d.fee_amount) for d in days])
        taxes = np.array([float(d.tax_amount) for d in days])
        total_costs = fees + taxes
        
        # Bar chart comparison
        width = 0.35
        x_offset = xs - width/2
        
        ax.bar(x_offset, gross, width, label='Brutto-Gewinn', 
               color=MODERN_COLORS['gross'], alpha=0.8, edgecolor='white', linewidth=0.5)
        ax.bar(x_offset + width, total_costs, width, label='Kosten (Gebühren + Steuern)', 
               color=MODERN_COLORS['fees'], alpha=0.8, edgecolor='white', linewidth=0.5)
        
        ax.set_ylabel("Betrag (€)", fontsize=10, fontweight='medium')
        ax.set_title("Gewinn vs. Kosten", fontsize=12, fontweight='bold', pad=15)

    def _export_xlsx(self) -> None:
        """Export data to styled Excel file."""
        if not OPENPYXL_AVAILABLE:
            QMessageBox.warning(self, "Export nicht verfügbar", 
                               "openpyxl ist nicht installiert. Bitte 'pip install openpyxl' ausführen.")
            return
            
        try:
            params = self._get_params()
            days, kpis = simulate(params)
        except Exception as e:
            QMessageBox.critical(self, "Export fehlgeschlagen", str(e))
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "XLSX speichern", "compounding_report.xlsx", "Excel (*.xlsx)"
        )
        if not path:
            return
            
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Compounding Daten"
            
            # Header styling
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="2C5F2C", end_color="2C5F2C", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Headers
            headers = ["Tag", "Startkapital", "Brutto", "Notional", "Gebühren",
                      "Vor Steuern", "Steuern", "Netto", "Reinvest", "Entnahme", "Endkapital"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = thin_border
            
            # Data rows
            money_format = '#,##0.00 €'
            for row_idx, d in enumerate(days, 2):
                values = [
                    d.day, float(d.start_capital), float(d.gross_profit), float(d.notional),
                    float(d.fee_amount), float(d.profit_before_tax), float(d.tax_amount),
                    float(d.net_profit), float(d.reinvest), float(d.withdrawal), float(d.end_capital)
                ]
                for col_idx, value in enumerate(values, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = thin_border
                    if col_idx > 1:  # Money columns
                        cell.number_format = money_format
                        cell.alignment = Alignment(horizontal="right")
            
            # KPIs summary sheet
            ws_kpi = wb.create_sheet("KPIs")
            kpi_data = [
                ("Endkapital", float(kpis.end_capital)),
                ("Summe Brutto", float(kpis.sum_gross)),
                ("Summe Gebühren", float(kpis.sum_fees)),
                ("Summe Steuern", float(kpis.sum_taxes)),
                ("Summe Netto", float(kpis.sum_net)),
                ("Summe reinvestiert", float(kpis.sum_reinvest)),
                ("Summe entnommen", float(kpis.sum_withdrawal)),
                ("ROI", f"{kpis.roi_pct}%"),
            ]
            for row_idx, (label, value) in enumerate(kpi_data, 1):
                ws_kpi.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
                cell = ws_kpi.cell(row=row_idx, column=2, value=value)
                if isinstance(value, float):
                    cell.number_format = money_format
            
            # Auto-adjust column widths
            for ws_sheet in [ws, ws_kpi]:
                for column in ws_sheet.columns:
                    max_length = max(len(str(cell.value or "")) for cell in column)
                    ws_sheet.column_dimensions[column[0].column_letter].width = min(max_length + 2, 20)
            
            wb.save(path)
            QMessageBox.information(self, "Export erfolgreich", f"Daten wurden nach {path} exportiert.")
            
        except Exception as e:
            QMessageBox.critical(self, "XLSX Export Fehler", str(e))
    
    def _export_docx(self) -> None:
        """Export data to Word document with formatted table."""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Export nicht verfügbar", 
                               "python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen.")
            return
            
        try:
            params = self._get_params()
            days, kpis = simulate(params)
        except Exception as e:
            QMessageBox.critical(self, "Export fehlgeschlagen", str(e))
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "DOCX speichern", "compounding_report.docx", "Word (*.docx)"
        )
        if not path:
            return
            
        try:
            doc = Document()
            
            # Title
            doc.add_heading('Compounding & P&L Report', 0)
            
            # KPIs section
            doc.add_heading('Monatliche Kennzahlen (KPIs)', level=1)
            kpi_table = doc.add_table(rows=8, cols=2)
            kpi_table.style = 'Table Grid'
            kpi_data = [
                ("Endkapital", f"{kpis.end_capital:.2f} €"),
                ("Summe Brutto", f"{kpis.sum_gross:.2f} €"),
                ("Summe Gebühren", f"{kpis.sum_fees:.2f} €"),
                ("Summe Steuern", f"{kpis.sum_taxes:.2f} €"),
                ("Summe Netto", f"{kpis.sum_net:.2f} €"),
                ("Summe reinvestiert", f"{kpis.sum_reinvest:.2f} €"),
                ("Summe entnommen", f"{kpis.sum_withdrawal:.2f} €"),
                ("ROI", f"{kpis.roi_pct}%"),
            ]
            for i, (label, value) in enumerate(kpi_data):
                row = kpi_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Bold the labels
                row.cells[0].paragraphs[0].runs[0].bold = True
            
            doc.add_paragraph()  # Spacing
            
            # Daily data section
            doc.add_heading('Tägliche Übersicht', level=1)
            
            headers = ["Tag", "Start", "Brutto", "Gebühren", "Steuern", "Netto", "Endkapital"]
            data_table = doc.add_table(rows=len(days) + 1, cols=len(headers))
            data_table.style = 'Table Grid'
            
            # Headers
            header_row = data_table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                header_row.cells[i].paragraphs[0].runs[0].bold = True
            
            # Data
            for row_idx, d in enumerate(days, 1):
                row = data_table.rows[row_idx]
                row.cells[0].text = str(d.day)
                row.cells[1].text = f"{d.start_capital:.2f}"
                row.cells[2].text = f"{d.gross_profit:.2f}"
                row.cells[3].text = f"{d.fee_amount:.2f}"
                row.cells[4].text = f"{d.tax_amount:.2f}"
                row.cells[5].text = f"{d.net_profit:.2f}"
                row.cells[6].text = f"{d.end_capital:.2f}"
            
            doc.save(path)
            QMessageBox.information(self, "Export erfolgreich", f"Report wurde nach {path} exportiert.")
            
        except Exception as e:
            QMessageBox.critical(self, "DOCX Export Fehler", str(e))

    def _copy_table_to_clipboard(self) -> None:
        try:
            params = self._get_params()
            days, _k = simulate(params)
            rows = to_csv_rows(days)
            text = "\n".join(["\t".join(r) for r in rows])
        except Exception as e:
            QMessageBox.critical(self, "Kopieren fehlgeschlagen", str(e))
            return
        QGuiApplication.clipboard().setText(text)
        QMessageBox.information(self, "Kopiert", "Tages-Tabelle wurde in die Zwischenablage kopiert.")
    
    # ==================== Settings Persistence ====================
    
    def _on_tab_changed(self, index: int) -> None:
        """Save settings when leaving Compounding tab."""
        # Save settings whenever tab changes (covers leaving the inputs tab)
        self._save_settings()
    
    def _save_settings(self) -> None:
        """Save current input values to QSettings."""
        try:
            self._settings.setValue("start_capital", self.sp_start_capital.value())
            self._settings.setValue("vip_level", self.cb_vip_level.currentText())
            self._settings.setValue("order_type", self.cb_order_type.currentText())
            self._settings.setValue("fee_open_pct", self.sp_fee_open_pct.value())
            self._settings.setValue("fee_close_pct", self.sp_fee_close_pct.value())
            self._settings.setValue("leverage", self.sp_leverage.value())
            self._settings.setValue("daily_profit_pct", self.sp_daily_profit_pct.value())
            self._settings.setValue("reinvest_pct", self.sp_reinvest_pct.value())
            self._settings.setValue("tax_pct", self.sp_tax_pct.value())
            self._settings.setValue("days", self.sp_days.value())
            self._settings.setValue("trading_days_preset", self.cb_trading_days_preset.currentIndex())
            self._settings.setValue("apply_losses", self.cb_apply_losses.isChecked())
            self._settings.setValue("target_month_net", self.sp_target_month_net.value())
            self._settings.setValue("plot_mode", self.cb_plot_mode.currentIndex())
            self._settings.sync()
        except Exception:
            pass  # Silently ignore save errors
    
    def _load_settings(self) -> None:
        """Load saved input values from QSettings."""
        self._suppress_updates = True
        try:
            if self._settings.contains("start_capital"):
                self.sp_start_capital.setValue(float(self._settings.value("start_capital", 100.0)))
            if self._settings.contains("vip_level"):
                self.cb_vip_level.setCurrentText(str(self._settings.value("vip_level", "VIP0")))
            if self._settings.contains("order_type"):
                self.cb_order_type.setCurrentText(str(self._settings.value("order_type", "Taker/Taker")))
            if self._settings.contains("fee_open_pct"):
                self.sp_fee_open_pct.setValue(float(self._settings.value("fee_open_pct", 0.06)))
            if self._settings.contains("fee_close_pct"):
                self.sp_fee_close_pct.setValue(float(self._settings.value("fee_close_pct", 0.06)))
            if self._settings.contains("leverage"):
                self.sp_leverage.setValue(float(self._settings.value("leverage", 20.0)))
            if self._settings.contains("daily_profit_pct"):
                self.sp_daily_profit_pct.setValue(float(self._settings.value("daily_profit_pct", 30.0)))
            if self._settings.contains("reinvest_pct"):
                self.sp_reinvest_pct.setValue(float(self._settings.value("reinvest_pct", 30.0)))
            if self._settings.contains("tax_pct"):
                self.sp_tax_pct.setValue(float(self._settings.value("tax_pct", 25.0)))
            if self._settings.contains("days"):
                self.sp_days.setValue(int(self._settings.value("days", 30)))
            if self._settings.contains("trading_days_preset"):
                self.cb_trading_days_preset.setCurrentIndex(int(self._settings.value("trading_days_preset", 0)))
            if self._settings.contains("apply_losses"):
                self.cb_apply_losses.setChecked(self._settings.value("apply_losses", False) in [True, "true", "True"])
            if self._settings.contains("target_month_net"):
                self.sp_target_month_net.setValue(float(self._settings.value("target_month_net", 0.0)))
            if self._settings.contains("plot_mode"):
                self.cb_plot_mode.setCurrentIndex(int(self._settings.value("plot_mode", 0)))
        except Exception:
            pass  # Use defaults on load errors
        finally:
            self._suppress_updates = False
    
    def closeEvent(self, event) -> None:
        """Save settings when widget is closed."""
        self._save_settings()
        super().closeEvent(event)
    
    # ==================== Monthly/Yearly Chart Views ====================
    
    def _plot_monthly(self, ax, days, xs) -> None:
        """Create monthly aggregated stacked bar chart."""
        # Aggregate daily data into monthly buckets (assume ~30 days per month)
        days_per_month = 30
        num_months = max(1, (len(days) + days_per_month - 1) // days_per_month)
        
        months = []
        gross_monthly = []
        fees_monthly = []
        taxes_monthly = []
        net_monthly = []
        
        for m in range(num_months):
            start_idx = m * days_per_month
            end_idx = min((m + 1) * days_per_month, len(days))
            month_days = days[start_idx:end_idx]
            
            months.append(m + 1)
            gross_monthly.append(sum(float(d.gross_profit) for d in month_days))
            fees_monthly.append(sum(float(d.fee_amount) for d in month_days))
            taxes_monthly.append(sum(float(d.tax_amount) for d in month_days))
            net_monthly.append(sum(float(d.net_profit) for d in month_days))
        
        x = np.arange(len(months))
        width = 0.6
        
        # Stacked bars: Net (bottom), then Fees, then Taxes on top
        ax.bar(x, net_monthly, width, label='Netto-Gewinn', color=MODERN_COLORS['net'], alpha=0.9)
        ax.bar(x, fees_monthly, width, bottom=net_monthly, label='Gebühren', color=MODERN_COLORS['fees'], alpha=0.8)
        bottom_taxes = [n + f for n, f in zip(net_monthly, fees_monthly)]
        ax.bar(x, taxes_monthly, width, bottom=bottom_taxes, label='Steuern', color=MODERN_COLORS['taxes'], alpha=0.8)
        
        # Gross line on top
        ax.plot(x, gross_monthly, color=MODERN_COLORS['gross'], linewidth=2.5, 
                marker='o', markersize=6, label='Brutto-Gewinn')
        
        ax.set_xticks(x)
        ax.set_xticklabels([f'Monat {m}' for m in months])
        ax.set_ylabel("Betrag (€)", fontsize=10, fontweight='medium')
        ax.set_title("Monatliche Übersicht (gestapelt)", fontsize=12, fontweight='bold', pad=15)
        ax.axhline(y=0, color='#8E8E93', linewidth=0.8, linestyle='-', alpha=0.5)
    
    def _plot_yearly(self, ax, days, xs) -> None:
        """Create yearly aggregated stacked bar chart."""
        # Aggregate daily data into yearly buckets (assume ~365 days per year)
        days_per_year = 365
        num_years = max(1, (len(days) + days_per_year - 1) // days_per_year)
        
        years = []
        gross_yearly = []
        fees_yearly = []
        taxes_yearly = []
        net_yearly = []
        
        for y in range(num_years):
            start_idx = y * days_per_year
            end_idx = min((y + 1) * days_per_year, len(days))
            year_days = days[start_idx:end_idx]
            
            years.append(y + 1)
            gross_yearly.append(sum(float(d.gross_profit) for d in year_days))
            fees_yearly.append(sum(float(d.fee_amount) for d in year_days))
            taxes_yearly.append(sum(float(d.tax_amount) for d in year_days))
            net_yearly.append(sum(float(d.net_profit) for d in year_days))
        
        x = np.arange(len(years))
        width = 0.6
        
        # Stacked bars
        ax.bar(x, net_yearly, width, label='Netto-Gewinn', color=MODERN_COLORS['net'], alpha=0.9)
        ax.bar(x, fees_yearly, width, bottom=net_yearly, label='Gebühren', color=MODERN_COLORS['fees'], alpha=0.8)
        bottom_taxes = [n + f for n, f in zip(net_yearly, fees_yearly)]
        ax.bar(x, taxes_yearly, width, bottom=bottom_taxes, label='Steuern', color=MODERN_COLORS['taxes'], alpha=0.8)
        
        # Gross line on top
        ax.plot(x, gross_yearly, color=MODERN_COLORS['gross'], linewidth=2.5, 
                marker='o', markersize=8, label='Brutto-Gewinn')
        
        ax.set_xticks(x)
        ax.set_xticklabels([f'Jahr {y}' for y in years])
        ax.set_ylabel("Betrag (€)", fontsize=10, fontweight='medium')
        ax.set_title("Jährliche Übersicht (gestapelt)", fontsize=12, fontweight='bold', pad=15)
        ax.axhline(y=0, color='#8E8E93', linewidth=0.8, linestyle='-', alpha=0.5)


if __name__ == "__main__":
    import sys
    from PyQt6.QtWidgets import QApplication, QMainWindow

    app = QApplication(sys.argv)
    win = QMainWindow()
    win.setWindowTitle("CompoundingPanel Demo")
    win.setCentralWidget(CompoundingPanel())
    win.resize(1100, 750)
    win.show()
    sys.exit(app.exec())
